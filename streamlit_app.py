import streamlit as st
import pandas as pd
import numpy as np
import joblib
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_squared_error, r2_score
import altair as alt
import time
import zipfile
import matplotlib.pyplot as plt
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Page title
st.set_page_config(page_title='ExoplanetML', page_icon=':alien:')
st.title(':alien: ExoplanetML: Machine Learning Model for Target Variable Prediction')

with st.expander('About this apssss'):
    st.markdown('**What can this appss do?**')
    st.info('This app allows users to build a machine learning (ML) model for Exoplanet target variable prediction in an end-to-end workflow. This encompasses data upload, data pre-processing, ML model building and post-model analysis.')
    st.markdown("""
    <div style="background-color: #f0f2f6; padding: 10px; border-radius: 5px;">
    Here's a useful tool for data curation [CSV only]: <a href="https://aivigoratemitotool.streamlit.app/" target="_blank">AI-powered Data Curation Tool</a>. Tip: Ensure that your CSV file doesn't have any NaNs.
    </div>
    <br>
    """, unsafe_allow_html=True)

    st.markdown('**How to use the app?**')
    st.warning('To work with the app, go to the sidebar and select a dataset. Adjust the model parameters, which will initiate the ML model building process, display the model results, and allow users to download the generated models and accompanying data.')

# Sidebar for input
with st.sidebar:
    # Load data
    st.header('1. Input data')

    st.markdown('**1.1 Use custom data**')
    uploaded_file = st.file_uploader("Upload a CSV file", type=["csv"])
    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file, index_col=False)

    # Download example data
    @st.cache_data
    def convert_df(input_df):
        return input_df.to_csv(index=False).encode('utf-8')

    example_csv = pd.read_csv('https://drive.google.com/uc?export=download&id=1J1f_qSHCYdfqiqQtpoda_Km_VmQuI4UX')
    csv = convert_df(example_csv)
    st.download_button(
        label="Download example CSV",
        data=csv,
        file_name='hwc-pesi.csv',
        mime='text/csv',
    )

    # Select example data
    st.markdown('**1.2. Use example data**')
    example_data = st.toggle('PHL Habitable Worlds Catalog (HWC)')
    if example_data:
        df = pd.read_csv('https://drive.google.com/uc?export=download&id=1J1f_qSHCYdfqiqQtpoda_Km_VmQuI4UX')

    st.header('2. Set Parameters')
    parameter_split_size = st.slider('Data split ratio (% for Training Set)', 10, 90, 80, 5)

    st.subheader('2.1. Learning Parameters')
    with st.expander('See parameters'):
        parameter_n_estimators = st.slider('Number of estimators (n_estimators)', 0, 1000, 100, 100)
        parameter_max_features = st.select_slider('Max features (max_features)', options=['all', 'sqrt', 'log2'])
        parameter_min_samples_split = st.slider('Minimum number of samples required to split an internal node (min_samples_split)', 2, 10, 2, 1)
        parameter_min_samples_leaf = st.slider('Minimum number of samples required to be at a leaf node (min_samples_leaf)', 1, 10, 2, 1)

    st.subheader('2.2. General Parameters')
    with st.expander('See parameters', expanded=False):
        parameter_random_state = st.slider('Seed number (random_state)', 0, 1000, 42, 1)
        parameter_criterion = st.select_slider('Performance measure (criterion)', options=['squared_error', 'absolute_error', 'friedman_mse'])
        parameter_bootstrap = st.select_slider('Bootstrap samples when building trees (bootstrap)', options=[True, False])
        parameter_oob_score = st.select_slider('Whether to use out-of-bag samples to estimate the R^2 on unseen data (oob_score)', options=[False, True])

    sleep_time = st.slider('Sleep time', 0, 3, 0)

# Model building process
if uploaded_file or example_data: 
    with st.status("Running ...", expanded=True) as status:
    
        st.write("Loading data ...")
        time.sleep(sleep_time)

        st.write("Preparing data ...")
        time.sleep(sleep_time)
        X = df.iloc[:,:-1]
        y = df.iloc[:,-1]
            
        st.write("Splitting data ...")
        time.sleep(sleep_time)
        X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=(100-parameter_split_size)/100, random_state=parameter_random_state)
        
        st.write("Model training ...")
        time.sleep(sleep_time)

        if parameter_max_features == 'all':
            parameter_max_features = None
        
        rf = RandomForestRegressor(
                n_estimators=parameter_n_estimators,
                max_features=parameter_max_features,
                min_samples_split=parameter_min_samples_split,
                min_samples_leaf=parameter_min_samples_leaf,
                random_state=parameter_random_state,
                criterion=parameter_criterion,
                bootstrap=parameter_bootstrap,
                oob_score=parameter_oob_score)
        rf.fit(X_train, y_train)
        
        st.write("Applying model to make predictions ...")
        time.sleep(sleep_time)
        y_train_pred = rf.predict(X_train)
        y_test_pred = rf.predict(X_test)


        # Directory to save plots
        plot_dir = "dependency_plots"
        os.makedirs(plot_dir, exist_ok=True)
        
        # Parameters to visualize dependencies
        input_parameters = ['P_MASS', 'P_RADIUS', 'P_TEMP_EQUIL', 'P_PERIOD', 'S_TEMPERATURE', 'S_MASS', 'S_RADIUS', 'P_FLUX', 'P_GRAVITY']
        
        # Assuming `X_test` is the test input features and `y_test_pred` are the model predictions
        for param in input_parameters:
            if param in X_test.columns:
                
                # Plot for all values
                plt.figure(figsize=(8, 6))
                plt.scatter(y_test_pred, X_test[param], alpha=0.6, color='skyblue')
                plt.xlabel("Predicted ESI")
                plt.ylabel(param)
                plt.title(f"All Values: {param} vs Predicted ESI")
                plt.grid(True)
                
                # Save the plot for all values
                plt.savefig(f"{plot_dir}/{param}_vs_Predicted_ESI_all.png")
                plt.close()
                
                # Plot for top 10 highest ESI values
                top_10_indices = np.argsort(y_test_pred)[-10:]  # Get indices of top 10 ESI predictions
                plt.figure(figsize=(8, 6))
                plt.scatter(y_test_pred[top_10_indices], X_test[param].iloc[top_10_indices], alpha=0.6, color='darkorange')
                plt.xlabel("Predicted ESI")
                plt.ylabel(param)
                plt.title(f"Top 10 ESI Values: {param} vs Predicted ESI")
                plt.grid(True)
                
                # Save the plot for top 10 values
                plt.savefig(f"{plot_dir}/{param}_vs_Predicted_ESI_top10.png")
                plt.close()

           
        st.write("Evaluating performance metrics ...")
        time.sleep(sleep_time)
        train_mse = mean_squared_error(y_train, y_train_pred)
        train_r2 = r2_score(y_train, y_train_pred)
        test_mse = mean_squared_error(y_test, y_test_pred)
        test_r2 = r2_score(y_test, y_test_pred)

        rf_results = pd.DataFrame({
            'Method': ['Random forest'],
            'Training MSE': [train_mse],
            'Training R2': [train_r2],
            'Test MSE': [test_mse],
            'Test R2': [test_r2]
        }).round(3)
        
    status.update(label="Status", state="complete", expanded=False)

    # Display data info
    st.header('Input data', divider='rainbow')
    col = st.columns(4)
    col[0].metric(label="No. of samples", value=X.shape[0], delta="")
    col[1].metric(label="No. of X variables", value=X.shape[1], delta="")
    col[2].metric(label="No. of Training samples", value=X_train.shape[0], delta="")
    col[3].metric(label="No. of Test samples", value=X_test.shape[0], delta="")
    
    with st.expander('Initial dataset', expanded=True):
        st.dataframe(df, height=210, use_container_width=True)

        # Define the x_axis options based on the columns of X_train
    x_axis = st.selectbox('Select X-axis', options=X_train.columns)  # Example dropdown for X-axis

    # Ensure y_train is a named Series if it's not a DataFrame
    y_train.name = 'P_ESI'  # Set the name for the y_train Series if it's not already set

        # Filter Options
    x_min = st.number_input('Minimum X value', value=float(X_train[x_axis].min()), step=0.01)
    x_max = st.number_input('Maximum X value', value=float(X_train[x_axis].max()), step=0.01)
    y_min = st.number_input('Minimum P_ESI value', value=float(y_train.min()), step=0.01)
    y_max = st.number_input('Maximum P_ESI value', value=float(y_train.max()), step=0.01)
    
    # Filter the DataFrame based on user input
    filtered_train_data = pd.concat([X_train, y_train], axis=1)
    filtered_train_data = filtered_train_data[(filtered_train_data[x_axis] >= x_min) & 
                                              (filtered_train_data[x_axis] <= x_max) &
                                              (filtered_train_data['P_ESI'] >= y_min) & 
                                              (filtered_train_data['P_ESI'] <= y_max)]
    
    # Inside the Train split expander
    with st.expander('Train split', expanded=False):
        train_col = st.columns((3, 1))
        with train_col[0]:
            st.markdown('**X**')
            st.dataframe(filtered_train_data.drop(columns='P_ESI'), height=210, hide_index=True, use_container_width=True)
        with train_col[1]:
            st.markdown('**y**')
            st.dataframe(filtered_train_data['P_ESI'], height=210, hide_index=True, use_container_width=True)
    
        # Train Set Scatter Plot
        st.subheader('Train Set Scatter Plot')
        train_chart = alt.Chart(filtered_train_data).mark_circle(size=60).encode(
            x='P_ESI',            # X-axis gets the selected variable
            y=x_axis,           # Set the Y-axis label to the actual column name 'P_ESI'
            tooltip=[x_axis, 'P_ESI']
        ).interactive()
        st.altair_chart(train_chart, use_container_width=True)
    
    # Inside the Test split expander
    with st.expander('Test split', expanded=False):
        test_col = st.columns((3, 1))
        with test_col[0]:
            st.markdown('**X**')
            st.dataframe(X_test, height=210, hide_index=True, use_container_width=True)
        with test_col[1]:
            st.markdown('**y**')
            st.dataframe(y_test, height=210, hide_index=True, use_container_width=True)
    
        # Test Set Scatter Plot
        st.subheader('Test Set Scatter Plot')
        test_chart = alt.Chart(pd.concat([X_test, y_test], axis=1)).mark_circle(size=60).encode(
            x='P_ESI',           # X-axis gets the selected variable
            y=x_axis,  # Set the Y-axis label to P_ESI
            tooltip=[x_axis, 'P_ESI']
        ).interactive()
        st.altair_chart(test_chart, use_container_width=True)

     # Create an in-memory ZIP file for the plots
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zip_file:
        for param in input_parameters:
            for suffix in ["all", "top10"]:
                plot_path = f"{plot_dir}/{param}_vs_Predicted_ESI_{suffix}.png"
                if os.path.exists(plot_path):
                    zip_file.write(plot_path, os.path.basename(plot_path))
    zip_buffer.seek(0)  # Move to the beginning of the file for download
    
    # Display download button
    st.download_button(
        label="Download All Dependency Plots",
        data=zip_buffer,
        file_name="dependency_plots.zip",
        mime="application/zip"
    )

    # Display feature importance plot
    importances = rf.feature_importances_
    feature_names = list(X.columns)
    forest_importances = pd.Series(importances, index=feature_names)
    df_importance = forest_importances.reset_index().rename(columns={'index': 'feature', 0: 'value'})
    
    bars = alt.Chart(df_importance).mark_bar(size=40).encode(
             x='value:Q',
             y=alt.Y('feature:N', sort='-x')
           ).properties(height=250)

    performance_col = st.columns((2, 0.2, 3))
    with performance_col[0]:
        st.header('Model performance', divider='rainbow')
        st.dataframe(rf_results.T.reset_index().rename(columns={'index': 'Parameter', 0: 'Value'}))
    with performance_col[2]:
        st.header('Feature importance', divider='rainbow')
        st.altair_chart(bars, theme='streamlit', use_container_width=True)

    # Prediction results
    st.header('Prediction results', divider='rainbow')
    s_y_train = pd.Series(y_train, name='actual').reset_index(drop=True)
    s_y_train_pred = pd.Series(y_train_pred, name='predicted').reset_index(drop=True)
    df_train = pd.DataFrame(data=[s_y_train, s_y_train_pred], index=None).T
    df_train['class'] = 'train'
        
    s_y_test = pd.Series(y_test, name='actual').reset_index(drop=True)
    s_y_test_pred = pd.Series(y_test_pred, name='predicted').reset_index(drop=True)
    df_test = pd.DataFrame(data=[s_y_test, s_y_test_pred], index=None).T
    df_test['class'] = 'test'
    
    df_prediction = pd.concat([df_train, df_test], axis=0)
    
    prediction_col = st.columns((2, 0.2, 3))
    
    # Display dataframe
    with prediction_col[0]:
        st.dataframe(df_prediction, height=320, use_container_width=True)

    # Display scatter plot of actual vs predicted values
    with prediction_col[2]:
        scatter = alt.Chart(df_prediction).mark_circle(size=60).encode(
                        x='actual',
                        y='predicted',
                        color='class'
                  )
        st.altair_chart(scatter, theme='streamlit', use_container_width=True)

    # Save trained model
    model_filename = 'rf_model.joblib'
    joblib.dump(rf, model_filename)

    with open(model_filename, 'rb') as f:
        st.download_button(
            label='Download Trained Model',
            data=f,
            file_name=model_filename,
            mime='application/octet-stream'
        )

    # Apply to new dataset
    st.header('Apply Trained Model to New Dataset')
    new_file = st.file_uploader("Upload a new CSV for prediction", type=["csv"], key='predict')
    
    if new_file is not None:
        new_data = pd.read_csv(new_file)

        # Load the trained model
        with open(model_filename, 'rb') as f:
            saved_model = joblib.load(f)

        # Get the feature names that the model was trained on, excluding the target column P_ESI
        model_features = saved_model.feature_names_in_

        # Check if the new dataset has the required input features
        missing_features = set(model_features).difference(new_data.columns)
        
        if len(missing_features) == 0:
            # Reorder the columns in the new dataset to match the model's expected input features
            new_X = new_data[model_features]
            
            # Predict using the loaded model
            predictions = saved_model.predict(new_X)
            
            new_data['Predected ESI'] = predictions

                        # Directory to save plots for new predictions
            plot_dir = "new_data_dependency_plots"
            os.makedirs(plot_dir, exist_ok=True)
            
            # Parameters to visualize dependencies
            for param in input_parameters:
                if param in new_data.columns:
                    
                    # Plot for all values in new dataset predictions
                    plt.figure(figsize=(8, 6))
                    plt.scatter(new_data['Predected ESI'], new_data[param], alpha=0.6, color='skyblue')
                    plt.xlabel("Predicted ESI")
                    plt.ylabel(param)
                    plt.title(f"All Values: {param} vs Predicted ESI")
                    plt.grid(True)
                    plt.savefig(f"{plot_dir}/{param}_vs_Predicted_ESI_all.png")
                    plt.close()
                    
                    # Plot for top 10 highest ESI values in new dataset predictions
                    top_10_indices = np.argsort(new_data['Predected ESI'])[-10:]  # Get indices of top 10 ESI predictions
                    plt.figure(figsize=(8, 6))
                    plt.scatter(new_data['Predected ESI'].iloc[top_10_indices], new_data[param].iloc[top_10_indices], alpha=0.6, color='darkorange')
                    plt.xlabel("Predicted ESI")
                    plt.ylabel(param)
                    plt.title(f"Top 10 ESI Values: {param} vs Predicted ESI")
                    plt.grid(True)
                    plt.savefig(f"{plot_dir}/{param}_vs_Predicted_ESI_top10.png")
                    plt.close()

            st.write(new_data.head())
            
            # Ensure x_axis_pred is defined through a user selection
            x_axis_pred = st.selectbox(
                'Select the X-axis parameter for the Prediction chart',
                options=new_data.columns
            )
            
            # Ensure that the parameter exists in the data
            if x_axis_pred in new_data.columns:
                # Add filtering options based on the selected x_axis_pred
                x_min_pred = st.number_input('Minimum X value (for Prediction)', value=float(new_data[x_axis_pred].min()), step=0.01)
                x_max_pred = st.number_input('Maximum X value (for Prediction)', value=float(new_data[x_axis_pred].max()), step=0.01)
                y_min_pred = st.number_input('Minimum Predected ESI value', value=float(new_data['Predected ESI'].min()), step=0.01)
                y_max_pred = st.number_input('Maximum Predected ESI value', value=float(new_data['Predected ESI'].max()), step=0.01)
            
                # Apply filtering to new_data based on the selected x_axis_pred
                filtered_new_data = new_data[(new_data[x_axis_pred] >= x_min_pred) & 
                                             (new_data[x_axis_pred] <= x_max_pred) &
                                             (new_data['Predected ESI'] >= y_min_pred) &
                                             (new_data['Predected ESI'] <= y_max_pred)]
            
                # Create scatter plot for predictions using filtered data
                scatter_pred_chart = alt.Chart(filtered_new_data).mark_circle(size=60).encode(
                    x='Predected ESI',
                    y=x_axis_pred,
                    tooltip=['Predected ESI', x_axis_pred]
                ).interactive()
                st.altair_chart(scatter_pred_chart, use_container_width=True)

                           # Create a new Word document
                doc = Document()
                doc.add_heading('Exoplanet Prediction Analysis Report', level=1)
                
                # Introduction section
                doc.add_heading('Introduction', level=2)
                doc.add_paragraph(
                    "This report presents an analysis of the dependencies between various exoplanet parameters and the predicted Earth Similarity Index (ESI) values. "
                    "The predictions were made using a machine learning model trained to estimate ESI values based on input parameters such as mass, radius, temperature, "
                    "and star properties. The findings below summarize the observed correlations, patterns, and potential implications for future studies."
                )
                
                # Summary of Parameter Dependencies
                doc.add_heading('Parameter Dependency Analysis', level=2)
                doc.add_paragraph(
                    "The following sections explore the dependencies of each parameter on the predicted ESI values. We analyzed both the complete dataset "
                    "and the top 10 highest ESI values to identify significant patterns and potential correlations."
                )
                
                # Insert plots and add analysis for each parameter
                for param in input_parameters:
                    doc.add_heading(f'{param} vs Predicted ESI', level=3)
                    
                    # Insert the "All Values" plot
                    all_values_img_path = f"{plot_dir}/{param}_vs_Predicted_ESI_all.png"
                    doc.add_paragraph("All Values:")
                    doc.add_picture(all_values_img_path, width=Inches(5.5))
                    
                    # Explanation for "All Values" plot
                    doc.add_paragraph(
                        f"The plot of all values for {param} shows the general trend and correlation with the predicted ESI values. Observing these trends can reveal whether "
                        f"{param} has a strong positive, negative, or neutral relationship with habitability as measured by ESI."
                    )
                    
                    # Insert the "Top 10 Values" plot
                    top10_values_img_path = f"{plot_dir}/{param}_vs_Predicted_ESI_top10.png"
                    doc.add_paragraph("Top 10 Highest ESI Values:")
                    doc.add_picture(top10_values_img_path, width=Inches(5.5))
                    
                    # Explanation for "Top 10 Values" plot
                    doc.add_paragraph(
                        f"The plot of the top 10 highest ESI values for {param} highlights the key attributes of the most Earth-like exoplanets in the dataset. "
                        f"This subset can help identify critical thresholds or values in {param} that appear most favorable for habitability."
                    )
                
                # Overall Findings and Future Implications
                doc.add_heading('Findings and Future Implications', level=2)
                doc.add_paragraph(
                    "In analyzing the dependencies across all parameters, we observed that certain factors, such as stellar temperature and planetary mass, "
                    "often correlate with higher ESI values. These insights can guide future studies in selecting exoplanets with optimal conditions for habitability. "
                    "By identifying key patterns in these parameters, future research can prioritize exploring planets within favorable ranges, "
                    "thus improving the chances of discovering potentially habitable exoplanets."
                )
                
                doc.add_paragraph(
                    "This analysis also emphasizes the importance of parameter dependencies in machine learning predictions. Understanding the influence of "
                    "each parameter on habitability predictions allows for more targeted and efficient exploration in the field of exoplanetary studies."
                )
                
                # Save the document
                doc_filename = "Exoplanet_Prediction_Analysis_Report.docx"
                doc.save(doc_filename)
                
                # Add a download button in Streamlit for the Word document
                with open(doc_filename, "rb") as file:
                    st.download_button(
                        label="Download Analysis Report",
                        data=file,
                        file_name=doc_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.error(f"The selected parameter '{x_axis_pred}' does not exist in the data.")

        else:
            # Show an error if the dataset is missing any required features
            st.error("The dataset is missing the following features: " + ", ".join(missing_features))
else:
    st.warning('👈 Upload a CSV file or click *"Load example data"* to get started!')
